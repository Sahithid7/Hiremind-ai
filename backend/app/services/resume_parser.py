import re
from dataclasses import dataclass
from pathlib import Path

import pdfplumber
from docx import Document
from PyPDF2 import PdfReader


TECH_SKILLS = {
    "python",
    "java",
    "javascript",
    "typescript",
    "react",
    "node",
    "fastapi",
    "django",
    "flask",
    "sql",
    "postgresql",
    "mysql",
    "mongodb",
    "aws",
    "azure",
    "gcp",
    "docker",
    "kubernetes",
    "terraform",
    "git",
    "github",
    "linux",
    "rest",
    "graphql",
    "redis",
    "celery",
    "pandas",
    "numpy",
    "machine learning",
    "openai",
    "ci/cd",
}

SKILL_DISPLAY_NAMES = {
    "aws": "AWS",
    "gcp": "GCP",
    "sql": "SQL",
    "ci/cd": "CI/CD",
    "fastapi": "FastAPI",
    "postgresql": "PostgreSQL",
    "javascript": "JavaScript",
    "typescript": "TypeScript",
    "github": "GitHub",
    "graphql": "GraphQL",
    "openai": "OpenAI",
    "rest": "REST",
}

SECTION_ALIASES = {
    "experience": {"experience", "work experience", "professional experience", "employment"},
    "education": {"education", "academic background"},
    "projects": {"projects", "technical projects", "personal projects"},
}


@dataclass
class ParsedResume:
    text: str
    skills: list[str]
    experience: list[dict]
    education: list[dict]
    projects: list[dict]


def parse_resume_file(path: Path) -> ParsedResume:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = extract_pdf_text(path)
    elif suffix == ".docx":
        text = extract_docx_text(path)
    else:
        raise ValueError("Unsupported resume file type.")

    normalized_text = normalize_text(text)
    return ParsedResume(
        text=normalized_text,
        skills=extract_skills(normalized_text),
        experience=extract_section_items(normalized_text, "experience"),
        education=extract_section_items(normalized_text, "education"),
        projects=extract_section_items(normalized_text, "projects"),
    )


def extract_pdf_text(path: Path) -> str:
    try:
        with pdfplumber.open(path) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
            text = "\n".join(pages).strip()
            if text:
                return text
    except Exception:
        pass

    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def extract_docx_text(path: Path) -> str:
    document = Document(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_text = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_text.append(" | ".join(cells))
    return "\n".join(paragraphs + table_text)


def normalize_text(text: str) -> str:
    lines = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


def extract_skills(text: str) -> list[str]:
    lowered = text.lower()
    matches = []
    for skill in sorted(TECH_SKILLS):
        pattern = rf"(?<![a-z0-9+#]){re.escape(skill)}(?![a-z0-9+#])"
        if re.search(pattern, lowered):
            matches.append(SKILL_DISPLAY_NAMES.get(skill, skill.title()))
    return matches


def extract_section_items(text: str, section: str) -> list[dict]:
    section_text = find_section_text(text, section)
    if not section_text:
        return []

    lines = [line.strip(" -•\t") for line in section_text.splitlines() if line.strip()]
    items = []
    current: list[str] = []

    for line in lines:
        looks_like_heading = bool(re.search(r"\b(20\d{2}|19\d{2}|present|intern|engineer|developer|university|college)\b", line, re.I))
        if looks_like_heading and current:
            items.append({"title": current[0], "details": current[1:4]})
            current = [line]
        else:
            current.append(line)

    if current:
        items.append({"title": current[0], "details": current[1:4]})

    return items[:6]


def find_section_text(text: str, section: str) -> str:
    aliases = SECTION_ALIASES[section]
    headings = []
    for alias in {value for values in SECTION_ALIASES.values() for value in values}:
        headings.append(re.escape(alias))

    section_pattern = "|".join(re.escape(alias) for alias in aliases)
    any_heading_pattern = "|".join(headings)
    match = re.search(
        rf"(?ims)^(?:{section_pattern})\s*$([\s\S]*?)(?=^(?:{any_heading_pattern})\s*$|\Z)",
        text,
    )
    return match.group(1).strip() if match else ""
